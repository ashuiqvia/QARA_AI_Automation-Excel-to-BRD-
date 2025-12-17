"""
Test script to verify that render_docx preserves all tables except Functional Requirements
"""
import sys
from io import BytesIO
from pathlib import Path
from docx import Document

# Fix Unicode encoding for Windows console
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

# Add parent directory to path to import from main.py
sys.path.insert(0, str(Path(__file__).parent))

from main import render_docx, parse_excel_to_requirements

def test_template_preservation():
    """Test that all tables are preserved except Functional Requirements"""
    
    # Load original template
    template_path = Path("templates/template.docx")
    with open(template_path, "rb") as f:
        template_bytes = f.read()
    
    original_doc = Document(BytesIO(template_bytes))
    original_table_count = len(original_doc.tables)
    
    print(f"Original template has {original_table_count} tables")
    
    # Create dummy requirements data
    dummy_requirements = [
        {
            "form": "Test Form",
            "requirements": [
                {
                    "req_id": "TEST_01.01",
                    "section": "Test Section",
                    "description": "Test Description",
                    "status": "Approved"
                }
            ]
        }
    ]
    
    # Render document
    output_stream = render_docx(template_bytes, dummy_requirements)
    output_bytes = output_stream.read()
    
    # Load rendered document
    rendered_doc = Document(BytesIO(output_bytes))
    rendered_table_count = len(rendered_doc.tables)
    
    print(f"Rendered document has {rendered_table_count} tables")
    
    # Check table preservation
    if rendered_table_count == original_table_count:
        print("✓ SUCCESS: All tables preserved!")
    else:
        print(f"✗ ERROR: Table count mismatch! Expected {original_table_count}, got {rendered_table_count}")
        return False
    
    # Verify specific tables exist
    table_names = []
    for i, table in enumerate(rendered_doc.tables):
        if len(table.rows) > 0:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            table_names.append(f"Table {i+1}: {headers[0] if headers else 'Empty'}")
    
    print("\nTables in rendered document:")
    for name in table_names:
        print(f"  {name}")
    
    # Check for DOCUMENT INFORMATION
    has_doc_info = False
    has_doc_history = False
    has_func_req = False
    
    for table in rendered_doc.tables:
        if len(table.rows) > 0:
            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if 'document name' in str(headers).lower():
                has_doc_info = True
            if 'revision number' in str(headers).lower():
                has_doc_history = True
            if 'requirement id' in str(headers).lower() and len(table.columns) == 4:
                has_func_req = True
    
    print("\n✓ DOCUMENT INFORMATION table preserved" if has_doc_info else "✗ DOCUMENT INFORMATION table MISSING")
    print("✓ DOCUMENT HISTORY table preserved" if has_doc_history else "✗ DOCUMENT HISTORY table MISSING")
    print("✓ FUNCTIONAL REQUIREMENTS table found" if has_func_req else "✗ FUNCTIONAL REQUIREMENTS table MISSING")
    
    if has_doc_info and has_doc_history and has_func_req:
        print("\n✅ ALL CHECKS PASSED!")
        return True
    else:
        print("\n❌ SOME CHECKS FAILED!")
        return False

if __name__ == "__main__":
    success = test_template_preservation()
    sys.exit(0 if success else 1)

