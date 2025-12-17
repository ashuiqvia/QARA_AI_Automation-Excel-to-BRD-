"""
Script to analyze the original and updated templates to see what's different
"""
from docx import Document
from pathlib import Path

def analyze_template(file_path):
    """Analyze a Word document and return its structure"""
    doc = Document(file_path)
    
    print(f"\n{'='*80}")
    print(f"Analyzing: {file_path}")
    print(f"{'='*80}")
    
    # Count paragraphs
    print(f"\nTotal Paragraphs: {len(doc.paragraphs)}")
    print("\nParagraphs (first 20):")
    for i, para in enumerate(doc.paragraphs[:20]):
        text = para.text.strip()[:80]
        if text:
            print(f"  [{i}] {text}")
    
    # Count tables
    print(f"\nTotal Tables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        print(f"\n  Table {i+1}:")
        print(f"    Rows: {len(table.rows)}")
        print(f"    Columns: {len(table.columns)}")
        if len(table.rows) > 0:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            print(f"    Headers: {headers}")
            
            # Check if this is DOCUMENT INFORMATION or DOCUMENT HISTORY
            if len(table.rows) > 0:
                first_row_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells]).lower()
                if 'document information' in first_row_text or 'document name' in first_row_text:
                    print(f"    *** This appears to be DOCUMENT INFORMATION table ***")
                elif 'document history' in first_row_text or 'revision' in first_row_text:
                    print(f"    *** This appears to be DOCUMENT HISTORY table ***")
                elif 'requirement' in first_row_text:
                    print(f"    *** This appears to be FUNCTIONAL REQUIREMENTS table ***")
    
    # Check sections
    print("\n\nSection Headings Found:")
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and (text.isupper() or 'DOCUMENT' in text.upper() or 'REQUIREMENT' in text.upper()):
            print(f"  - {text}")

if __name__ == "__main__":
    templates_dir = Path("templates")
    
    original = templates_dir / "Business Requirements Document_Orginal.docx"
    updated = templates_dir / "Business Requirements Document - updated (11).docx"
    
    if original.exists():
        analyze_template(original)
    else:
        print(f"Original template not found: {original}")
    
    if updated.exists():
        analyze_template(updated)
    else:
        print(f"Updated template not found: {updated}")

