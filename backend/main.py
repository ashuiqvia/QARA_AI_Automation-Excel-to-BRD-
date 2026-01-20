import logging
import re
from io import BytesIO
from typing import Optional
from pydantic import BaseModel, EmailStr

import pandas as pd # type: ignore 
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, status, Depends # type: ignore 
from fastapi.middleware.cors import CORSMiddleware # type: ignore 
from starlette.responses import StreamingResponse, JSONResponse # type: ignore 
from docx import Document # type: ignore
from docx.shared import Pt, RGBColor # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH # type: ignore
from docx.oxml import OxmlElement # type: ignore
from docx.oxml.ns import qn # type: ignore

# Import authentication and database modules
from auth import verify_password, get_password_hash, create_access_token, get_current_user
from database import init_database, create_user, get_user_by_username, get_user_by_email 

# ------------------------------------------------------------------------------
# App & CORS
# ------------------------------------------------------------------------------
app = FastAPI()

# ------------------------------------------------------------------------------
# Authentication Models
# ------------------------------------------------------------------------------
class UserRegister(BaseModel):
    username: str
    email: EmailStr
    password: str
    full_name: Optional[str] = None

class UserLogin(BaseModel):
    username: str
    password: str

class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    username: str

# ------------------------------------------------------------------------------
# Authentication Routes
# ------------------------------------------------------------------------------
@app.post("/api/auth/register", response_model=TokenResponse)
async def register(user_data: UserRegister):
    """Register a new user."""
    try:
        # Check if username already exists
        try:
            if get_user_by_username(user_data.username):
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Username already exists"
                )
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Database error checking username: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail=f"Database connection error. Please check your database configuration. Error: {str(e)}"
            )
        
        # Check if email already exists
        try:
            if get_user_by_email(user_data.email):
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Email already exists"
                )
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Database error checking email: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail=f"Database connection error. Please check your database configuration. Error: {str(e)}"
            )
        
        # Hash password
        password_hash = get_password_hash(user_data.password)
        
        # Create user
        try:
            success = create_user(
                username=user_data.username,
                email=user_data.email,
                password_hash=password_hash,
                full_name=user_data.full_name
            )
        except Exception as e:
            logger.error(f"Database error creating user: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail=f"Database connection error. Please check your database configuration. Error: {str(e)}"
            )
        
        if not success:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Failed to create user"
            )
        
        # Create access token
        access_token = create_access_token(data={"sub": user_data.username})
        
        return TokenResponse(
            access_token=access_token,
            username=user_data.username
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Unexpected error during registration")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Internal server error: {str(e)}"
        )

@app.post("/api/auth/login", response_model=TokenResponse)
async def login(credentials: UserLogin):
    """Login and get access token."""
    try:
        # Get user from database
        try:
            user = get_user_by_username(credentials.username)
        except Exception as e:
            logger.error(f"Database error during login: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail=f"Database connection error. Please check your database configuration. Error: {str(e)}"
            )
        
        if not user:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Incorrect username or password"
            )
        
        # Verify password
        if not verify_password(credentials.password, user["password_hash"]):
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Incorrect username or password"
            )
        
        # Create access token
        access_token = create_access_token(data={"sub": user["username"], "user_id": user["id"]})
        
        return TokenResponse(
            access_token=access_token,
            username=user["username"]
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Unexpected error during login")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Internal server error: {str(e)}"
        )

@app.get("/api/auth/me")
async def get_current_user_info(current_user: dict = Depends(get_current_user)):
    """Get current authenticated user information."""
    user = get_user_by_username(current_user["username"])
    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found"
        )
    
    return {
        "id": user["id"],
        "username": user["username"],
        "email": user["email"],
        "full_name": user["full_name"]
    }

# Health route so you can check server availability
@app.get("/health")
def health():
    return {"status": "ok"}

# Test endpoint to see parsed data structure
@app.post("/test-parse")
async def test_parse(excel: UploadFile = File(...), sheet_name: str | None = Form(None)):
    """Test endpoint to see how Excel is being parsed"""
    try:
        excel_bytes = await excel.read()
        requirements = parse_excel_to_requirements(excel_bytes, sheet_name=sheet_name, filter_mode="none")
        return {
            "total_groups": len(requirements),
            "groups": requirements,
            "sample_structure": requirements[0] if requirements else None
        }
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=400)

# CORS configuration - supports both development and production
import os
from dotenv import load_dotenv

load_dotenv()
ENVIRONMENT = os.getenv("ENVIRONMENT", "development")

if ENVIRONMENT == "production":
    # Production: Get allowed origins from environment or use default
    allowed_origins_str = os.getenv("ALLOWED_ORIGINS", "")
    if allowed_origins_str:
        allowed_origins = [origin.strip() for origin in allowed_origins_str.split(",")]
    else:
        # Default production origins (update these with your actual frontend URLs)
        allowed_origins = [
            "https://your-frontend.vercel.app",
            "https://your-frontend.netlify.app",
        ]
    allow_origin_regex = None
else:
    # Development: Allow localhost
    allowed_origins = [
        "http://localhost:5173",
        "http://localhost:5174",
        "http://127.0.0.1:5173",
        "http://127.0.0.1:5174",
    ]
    allow_origin_regex = r"http://(localhost|127\.0\.0\.1):\d+"

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_origin_regex=allow_origin_regex,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ------------------------------------------------------------------------------
# Logging (helpful to see what routes are loaded on startup)
# ------------------------------------------------------------------------------
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("brd-utility")

@app.on_event("startup")
async def on_startup():
    # Initialize database and create tables
    try:
        init_database()
        logger.info("Database initialized successfully")
    except Exception as e:
        logger.error(f"Database initialization failed: {str(e)}")
        logger.warning("Application will continue, but authentication may not work")
    
    logger.info("FastAPI started. Listing routes:")
    for r in app.routes:
        logger.info("Route loaded: %s %s", getattr(r, "methods", None), getattr(r, "path", None))


# ------------------------------------------------------------------------------
# Excel parsing helpers
# ------------------------------------------------------------------------------
def normalize_header(h: str) -> str:
    """Normalize Excel header: lowercase, collapse spaces."""
    normalized = re.sub(r"\s+", " ", str(h or "")).strip().lower()
    # Handle status column variations - normalize all status columns to "status *"
    if "status" in normalized:
        return "status *"
    return normalized

# Expected (case/space-insensitive) headers from your Excel:
EXPECTED_HEADERS = [
    "form",
    "req id#*",
    "section*",
    "description *",
    "status *",
]

def parse_excel_to_requirements(
    excel_bytes: bytes,
    sheet_name: Optional[str] = None,
    filter_mode: str = "none",
):
    """
    Read Excel, validate headers, and return a list of dicts with keys:
    req_id, section, description, status

    Option A fix:
    - If sheet_name is not provided, default to the FIRST sheet (index 0),
      so pandas returns a single DataFrame (not a dict of DataFrames).
    """
    # ✅ Option A: default to first sheet when sheet_name is None/empty
    target_sheet = sheet_name if sheet_name else 0
    df = pd.read_excel(BytesIO(excel_bytes), sheet_name=target_sheet, engine="openpyxl")

    # Normalize column names
    df.columns = [normalize_header(c) for c in df.columns]

    # Validate required columns exist
    missing = [c for c in EXPECTED_HEADERS if c not in df.columns]
    if missing:
        raise ValueError(f"Missing required Excel columns: {missing}")

    # Transform rows - create flat list with form info
    reqs = []
    current_form = None
    
    for _, row in df.iterrows():
        # Get Form (section heading) - keep track of current form
        form = str(row.get("form", "")).strip()
        if form:
            current_form = form
        
        # Get requirement ID
        req_id = str(row.get("req id#*", "")).strip()
        if not req_id:
            continue  # skip blank id rows

        item = {
            "req_id": req_id,
            "section": str(row.get("section*", "")).strip(),
            "description": str(row.get("description *", "")).strip(),
            "status": str(row.get("status *", "")).strip(),
            "form": current_form if current_form else "",  # Include form with each requirement
        }
        reqs.append(item)
    
    # Apply filtering if needed
    fm = (filter_mode or "none").lower()
    if fm == "final":
        reqs = [r for r in reqs if r["status"].lower() == "final"]
    elif fm == "final_or_approved":
        reqs = [r for r in reqs if r["status"].lower() in ("final", "approved")]
    
    # Also create grouped structure for templates that need it
    reqs_by_form = {}
    for req in reqs:
        form_name = req.get("form", "Other")
        if form_name not in reqs_by_form:
            reqs_by_form[form_name] = []
        reqs_by_form[form_name].append(req)
    
    # Return both flat list and grouped structure
    grouped_reqs = []
    for form_name, form_reqs in reqs_by_form.items():
        grouped_reqs.append({
            "form": form_name,
            "requirements": form_reqs
        })
    
    logger.info(f"Parsed {len(reqs)} requirements in {len(grouped_reqs)} form groups")
    
    # Return grouped structure for template
    return grouped_reqs


# ------------------------------------------------------------------------------
# Word rendering - Programmatic approach matching brd_updater.py logic
# ------------------------------------------------------------------------------

# Color constants matching brd_updater.py
TABLE_HEADER_COLOR = (0, 176, 240)  # RGB: Blue
FORM_HEADER_COLOR = (0, 176, 240)   # RGB: Blue

def _set_cell_background(cell, color: tuple):
    """Set cell background color - matching brd_updater.py logic"""
    cell_properties = cell._element.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    # Convert RGB to hex (e.g., (0, 176, 240) → '00B0F0')
    hex_color = '{:02X}{:02X}{:02X}'.format(color[0], color[1], color[2])
    shading.set(qn('w:fill'), hex_color)
    cell_properties.append(shading)

def _format_header_row(table):
    """Format table header row - matching brd_updater.py format_header_row()"""
    header_row = table.rows[0]  # First row
    for cell in header_row.cells:
        # Background color: Blue (RGB: 0, 176, 240)
        _set_cell_background(cell, TABLE_HEADER_COLOR)
        
        # Text formatting
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)  # White text

def _add_form_header(table, form_name: str):
    """Add form header row - matching brd_updater.py _add_form_header()"""
    row = table.add_row()
    merged_cell = row.cells[0]
    
    # Set text
    merged_cell.text = form_name
    
    # Background color: Blue
    _set_cell_background(merged_cell, FORM_HEADER_COLOR)
    
    # Text formatting
    for paragraph in merged_cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
    
    # Cell merging (all 4 columns merged into 1)
    if len(row.cells) >= 4:
        merged_cell.merge(row.cells[1])
        merged_cell.merge(row.cells[2])
        merged_cell.merge(row.cells[3])

def _add_requirement_row(table, req: dict):
    """Add requirement data row - matching brd_updater.py _add_requirement_row()"""
    row = table.add_row()
    
    # Set cell values
    row.cells[0].text = req.get("req_id", "")
    row.cells[1].text = req.get("section", "")
    row.cells[2].text = req.get("description", "")
    row.cells[3].text = req.get("status", "")
    
    # Formatting
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            # Font size
            for run in paragraph.runs:
                run.font.size = Pt(10)
            # Alignment
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # No background color (default white)

def render_docx(template_bytes: bytes, requirements: list):
    """
    Programmatically build Word document matching brd_updater.py logic.
    Requirements should be grouped structure: [{"form": "...", "requirements": [...]}, ...]
    
    IMPORTANT: This function ONLY updates the Functional Requirements table.
    All other content (DOCUMENT INFORMATION, DOCUMENT HISTORY, etc.) is preserved.
    """
    # Load template document - this preserves ALL content including all tables
    doc = Document(BytesIO(template_bytes))
    
    logger.info(f"Loaded template with {len(doc.tables)} tables and {len(doc.paragraphs)} paragraphs")
    
    # Find the Functional Requirements table specifically
    # Look for table with 4 columns and headers: Requirement ID, Section, Description, Status
    target_table = None
    expected_headers = ['requirement id', 'section', 'description', 'status']
    
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) > 0 and len(table.columns) == 4:
            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
            # Check if this matches the Functional Requirements table structure
            # Must have "requirement id" in first column and "section" in second
            if (len(headers) >= 4 and 
                'requirement id' in headers[0] and 
                'section' in headers[1] and
                'description' in headers[2] and
                'status' in headers[3]):
                target_table = table
                logger.info(f"Found Functional Requirements table at index {table_idx} with headers: {headers}")
                break
    
    if not target_table:
        logger.warning("Functional Requirements table not found. Searching for insertion point...")
        # Find where to insert (after "Functional Requirements" heading)
        insert_pos = None
        for i, para in enumerate(doc.paragraphs):
            if 'functional requirements' in para.text.lower():
                insert_pos = i + 1
                break
        
        if insert_pos:
            # Create table after the heading
            target_table = doc.add_table(rows=1, cols=4)
            # Set headers
            headers = ['Requirement ID', 'Section', 'Description', 'Status']
            for i, header in enumerate(headers):
                target_table.rows[0].cells[i].text = header
            logger.info("Created new Functional Requirements table")
        else:
            raise ValueError("Could not find 'Functional Requirements' section in template")
    else:
        # Clear existing data rows (keep header row only)
        # Remove rows from end to beginning to avoid index issues
        initial_row_count = len(target_table.rows)
        for i in range(len(target_table.rows) - 1, 0, -1):
            target_table._element.remove(target_table.rows[i]._element)
        logger.info(f"Cleared {initial_row_count - 1} existing data rows from Functional Requirements table")
    
    # Format header row
    _format_header_row(target_table)
    
    # Process requirements grouped by Form
    total_reqs = 0
    for group in requirements:
        if isinstance(group, dict) and "requirements" in group:
            form_name = group.get("form", "")
            form_reqs = group.get("requirements", [])
            
            if form_name and form_reqs:
                # Add form header (merged row)
                _add_form_header(target_table, form_name)
                
                # Add requirement rows
                for req in form_reqs:
                    _add_requirement_row(target_table, req)
                    total_reqs += 1
    
    logger.info(f"Rendered Functional Requirements table with {len(target_table.rows)} rows ({total_reqs} requirements)")
    logger.info(f"Document still has {len(doc.tables)} tables total (all other tables preserved)")
    
    # Save to BytesIO - this preserves ALL tables and content
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


# ------------------------------------------------------------------------------
# /generate endpoint
# ------------------------------------------------------------------------------
@app.post("/generate")
async def generate_brd(
    excel: UploadFile = File(..., description="Excel file with requirements"),
    template: UploadFile | None = File(None, description="Optional Word template; if absent, server template is used"),
    sheet_name: str | None = Form(None),
    filter_mode: str = Form("none"),  # options: "none" | "final" | "final_or_approved"
    current_user: dict = Depends(get_current_user),  # Require authentication
):
    """
    Accepts the Excel (and optional Word template), generates a BRD docx with
    the Functional Requirements table updated under Section 2, and returns the file.
    """
    try:
        excel_bytes = await excel.read()

        # Use uploaded template if provided; else use server-side template
        if template:
            template_bytes = await template.read()
        else:
            with open("templates/template.docx", "rb") as f:
                template_bytes = f.read()

        requirements = parse_excel_to_requirements(
            excel_bytes,
            sheet_name=sheet_name,
            filter_mode=filter_mode,
        )

        if not requirements:
            return JSONResponse(
                {"message": "No requirements matched with the selected filter."},
                status_code=422,
            )
        
        logger.info(f"Parsed {len(requirements)} form groups from Excel")
        for i, group in enumerate(requirements):
            logger.info(f"  Group {i+1}: Form='{group.get('form')}', Requirements={len(group.get('requirements', []))}")

        output_stream = render_docx(template_bytes, requirements)

        filename = "Business Requirements Document - updated.docx"
        return StreamingResponse(
            output_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except ValueError as ve:
        logger.exception("Validation error during generation")
        return JSONResponse({"error": str(ve)}, status_code=400)
    except Exception as e:
        logger.exception("Unexpected error during generation")
        return JSONResponse({"error": f"Internal server error: {str(e)}"}, status_code=500)
