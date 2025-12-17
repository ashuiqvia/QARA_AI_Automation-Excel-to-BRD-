"""
Script to verify the template structure and show what variables are available
"""
from docx import Document
import re

def verify_template():
    try:
        doc = Document('templates/template.docx')
        
        print("=" * 60)
        print("TEMPLATE VERIFICATION")
        print("=" * 60)
        
        # Check paragraphs
        print("\n1. PARAGRAPHS:")
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                print(f"   Para {i+1}: {text[:50]}")
                # Check for Jinja2 syntax
                if '{%' in text or '{{' in text:
                    print(f"      [OK] Contains Jinja2 syntax")
        
        # Check tables
        print(f"\n2. TABLES: Found {len(doc.tables)} table(s)")
        for table_idx, table in enumerate(doc.tables):
            print(f"\n   Table {table_idx + 1}:")
            print(f"   Rows: {len(table.rows)}, Columns: {len(table.columns)}")
            
            # Check header row
            if len(table.rows) > 0:
                header_row = table.rows[0]
                headers = [cell.text.strip() for cell in header_row.cells]
                print(f"   Headers: {headers}")
            
            # Check data rows for Jinja2
            if len(table.rows) > 1:
                print(f"\n   Data Rows (checking for Jinja2):")
                for row_idx, row in enumerate(table.rows[1:], start=2):
                    cells_text = [cell.text.strip() for cell in row.cells]
                    has_jinja = any('{{' in text or '{%' in text for text in cells_text)
                    status = 'HAS Jinja2' if has_jinja else 'NO Jinja2'
                    print(f"   Row {row_idx}: {status}")
                    if has_jinja:
                        for cell_idx, text in enumerate(cells_text):
                            if '{{' in text or '{%' in text:
                                print(f"      Cell {cell_idx+1}: {text}")
        
        # Search for Jinja2 patterns
        print(f"\n3. JINJA2 SYNTAX CHECK:")
        jinja_patterns = [
            r'\{\%\s*for\s+.*\s+in\s+.*\s*\%\}',
            r'\{\%\s*endfor\s*\%\}',
            r'\{\{\s*req\.req_id\s*\}\}',
            r'\{\{\s*req\.section\s*\}\}',
            r'\{\{\s*req\.description\s*\}\}',
            r'\{\{\s*req\.status\s*\}\}',
        ]
        
        found_patterns = []
        all_text = ""
        for para in doc.paragraphs:
            all_text += para.text + " "
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "
        
        for pattern in jinja_patterns:
            matches = re.findall(pattern, all_text, re.IGNORECASE)
            if matches:
                found_patterns.append(pattern)
                print(f"   [OK] Found: {pattern}")
        
        if not found_patterns:
            print("   [ERROR] NO Jinja2 syntax found in template!")
            print("\n   You need to add:")
            print("   - {% for req in requirements_flat %}")
            print("   - {{ req.req_id }}, {{ req.section }}, {{ req.description }}, {{ req.status }}")
            print("   - {% endfor %}")
        
        print("\n" + "=" * 60)
        print("VERIFICATION COMPLETE")
        print("=" * 60)
        
    except Exception as e:
        print(f"Error reading template: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    verify_template()

