"""
Script to create a simple Word template for testing
Run this to generate a basic template.docx file
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_template():
    doc = Document()
    
    # Add title
    title = doc.add_heading('Functional Requirements', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add paragraph break
    doc.add_paragraph()
    
    # Create table with 4 columns
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Requirement ID'
    header_cells[1].text = 'Section'
    header_cells[2].text = 'Description'
    header_cells[3].text = 'Status'
    
    # Make header bold and colored
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set header background color (blue)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '4472C4')  # Blue color
    for cell in header_cells:
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Data row with Jinja2 template syntax
    data_cells = table.rows[1].cells
    data_cells[0].text = '{% for req in requirements_flat %}'
    data_cells[1].text = '{{ req.req_id }}'
    data_cells[2].text = '{{ req.section }}'
    data_cells[3].text = '{{ req.description }}'
    
    # Add another row for status and endfor
    row = table.add_row()
    row.cells[0].text = '{{ req.status }}'
    row.cells[1].text = '{% endfor %}'
    
    # Save template
    doc.save('templates/template.docx')
    print("Template created successfully at templates/template.docx")
    print("\nTemplate structure:")
    print("- Header row with: Requirement ID | Section | Description | Status")
    print("- Data row with Jinja2: {% for req in requirements_flat %}")
    print("- Variables: {{ req.req_id }}, {{ req.section }}, {{ req.description }}, {{ req.status }}")
    print("- End loop: {% endfor %}")

if __name__ == "__main__":
    create_template()

