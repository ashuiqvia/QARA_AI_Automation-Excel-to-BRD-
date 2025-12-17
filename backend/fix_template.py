"""
Script to fix the template by ensuring {% endfor %} is present
"""
from docx import Document
import re

def fix_template():
    doc = Document('templates/template.docx')
    
    # Check if {% endfor %} exists
    has_endfor = False
    all_text = ""
    
    for para in doc.paragraphs:
        text = para.text
        all_text += text + " "
        if '{% endfor %}' in text or '{%endfor%}' in text:
            has_endfor = True
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                all_text += text + " "
                if '{% endfor %}' in text or '{%endfor%}' in text:
                    has_endfor = True
    
    print("Template Fix Check:")
    print("=" * 60)
    
    if '{% for req in requirements_flat %}' in all_text or '{%for req in requirements_flat%}' in all_text:
        print("[OK] Found: {% for req in requirements_flat %}")
    else:
        print("[ERROR] Missing: {% for req in requirements_flat %}")
    
    if has_endfor:
        print("[OK] Found: {% endfor %}")
    else:
        print("[ERROR] Missing: {% endfor %} - Adding it now...")
        # Add {% endfor %} after the table
        if doc.tables:
            # Add paragraph after last table
            endfor_para = doc.add_paragraph('{% endfor %}')
            endfor_para.runs[0].font.name = 'Courier New'
            endfor_para.runs[0].font.size = 10
            doc.save('templates/template.docx')
            print("[OK] Added {% endfor %} to template")
    
    # Verify all variables
    variables = ['req.req_id', 'req.section', 'req.description', 'req.status']
    print("\nVariable Check:")
    for var in variables:
        if f'{{{{ {var} }}}}' in all_text or f'{{{{{var}}}}}' in all_text:
            print(f"[OK] Found: {{ {var} }}")
        else:
            print(f"[ERROR] Missing: {{ {var} }}")
    
    print("\n" + "=" * 60)
    print("Template should be ready now!")
    print("=" * 60)

if __name__ == "__main__":
    fix_template()

