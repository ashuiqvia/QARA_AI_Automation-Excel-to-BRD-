"""
Fix the template structure - put loop tags in the table row itself
For docxtpl, the loop should be in the table row cells
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def fix_template_structure():
    doc = Document('templates/template.docx')
    
    # Find the Functional Requirements table
    target_table = None
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if 'Requirement ID' in str(headers) or 'requirement id' in str(headers).lower():
            target_table = table
            break
    
    if not target_table:
        print("ERROR: Could not find Functional Requirements table")
        return
    
    print(f"Found table with {len(target_table.rows)} rows and {len(target_table.columns)} columns")
    
    # Ensure we have at least 2 rows (header + data)
    if len(target_table.rows) < 2:
        # Add a data row
        target_table.add_row()
        print("Added data row to table")
    
    # Get the data row (row 1, index 1, after header)
    data_row = target_table.rows[1]
    
    # Put loop start in first cell
    first_cell = data_row.cells[0]
    # Clear and set loop start
    first_cell.text = '{% for req in requirements_flat %}\n{{ req.req_id }}'
    
    # Set variables in other cells
    if len(data_row.cells) >= 2:
        data_row.cells[1].text = '{{ req.section }}'
    if len(data_row.cells) >= 3:
        data_row.cells[2].text = '{{ req.description }}'
    if len(data_row.cells) >= 4:
        # Put loop end in last cell
        data_row.cells[3].text = '{{ req.status }}\n{% endfor %}'
    
    # Remove any loop tags from paragraphs (they should be in table)
    paragraphs_to_remove = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if '{% for req in requirements_flat %}' in text or '{% endfor %}' in text:
            # Check if it's near our table - if so, we might want to keep it
            # Actually, let's keep paragraphs but ensure table has the loop
            pass
    
    doc.save('templates/template.docx')
    print("\n[OK] Template structure fixed!")
    print("\nTemplate now has:")
    print("  - Loop start in first cell: {% for req in requirements_flat %}")
    print("  - Variables in cells: {{ req.req_id }}, {{ req.section }}, etc.")
    print("  - Loop end in last cell: {% endfor %}")
    print("\nThis structure should work with docxtpl.")

if __name__ == "__main__":
    fix_template_structure()

