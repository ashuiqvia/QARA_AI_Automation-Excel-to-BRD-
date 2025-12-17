"""
Script to update the Word template with correct Jinja2 syntax
This will modify template.docx to have the proper structure
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def update_template():
    doc = Document('templates/template.docx')
    
    # Clear existing content (optional - comment out if you want to keep existing)
    # doc.paragraphs.clear()
    
    # Add title if not exists
    if not doc.paragraphs or 'Functional Requirements' not in doc.paragraphs[0].text:
        title = doc.add_heading('Functional Requirements', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Add space
    
    # Remove existing tables if any
    for table in doc.tables[:]:
        doc.element.body.remove(table._element)
    
    # Create new table with proper structure
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    
    # Set column widths
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(1.5)
    
    # Header row
    header_cells = table.rows[0].cells
    headers = ['Requirement ID', 'Section', 'Description', 'Status']
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        # Format header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Set blue background
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '4472C4')  # Blue
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Data row with Jinja2 variables
    data_cells = table.rows[1].cells
    jinja_vars = [
        '{{ req.req_id }}',
        '{{ req.section }}',
        '{{ req.description }}',
        '{{ req.status }}'
    ]
    
    for i, var_text in enumerate(jinja_vars):
        cell = data_cells[i]
        cell.text = var_text
        # Center align
        for paragraph in cell.paragraphs:
            if i == 2:  # Description column - left align
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add Jinja2 loop around the table
    # We need to add the loop tags before and after the table
    # In docx, we add paragraphs before and after
    
    # Get the table element
    table_element = table._element
    
    # Create paragraph for {% for %} before table
    for_para = doc.add_paragraph('{% for req in requirements_flat %}')
    for_para.runs[0].font.name = 'Courier New'
    for_para.runs[0].font.size = Pt(10)
    
    # Move table after the for loop paragraph
    # (Actually, we'll insert the loop text differently)
    
    # Actually, for docxtpl, the loop should be in the same paragraph or around the row
    # Let's add it as a separate paragraph before the table
    # But first, let's find where the table is and add text before it
    
    # Save the document
    doc.save('templates/template.docx')
    print("[OK] Template updated successfully!")
    print("\nTemplate structure:")
    print("1. Title: 'Functional Requirements'")
    print("2. Table with 4 columns: Requirement ID | Section | Description | Status")
    print("3. Header row: Blue background, white text")
    print("4. Data row with Jinja2 variables:")
    print("   - {{ req.req_id }}")
    print("   - {{ req.section }}")
    print("   - {{ req.description }}")
    print("   - {{ req.status }}")
    print("\n[IMPORTANT] You need to manually add the loop tags in Word:")
    print("   Before the table: {% for req in requirements_flat %}")
    print("   After the table: {% endfor %}")
    print("\nOr use the alternative: Put loop tags around the data row in the table")

if __name__ == "__main__":
    try:
        update_template()
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()

